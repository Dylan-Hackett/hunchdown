"""
Account Locator docx: a master index of each subject's main account per platform.

A "main account" is a capture whose URL points to a profile/landing page rather
than a single post (detected by platforms.is_main_account_url).

For v0.1 we assume one subject per case, so the doc has one row per main-account
capture, grouped by platform in priority order.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from .parser import Capture
from .platforms import (
    PLATFORM_DISPLAY_NAMES,
    PLATFORM_ORDER,
    classify,
    extract_handle,
    is_main_account_url,
)


@dataclass
class AccountEntry:
    platform: str
    handle: str | None
    url: str
    capture_date_raw: str
    exhibit_doc_filename: str | None      # cross-reference, if any posts exist
    exhibit_count: int                    # exhibits in the cross-referenced doc


def collect_main_accounts(captures: list[Capture]) -> list[Capture]:
    out: list[Capture] = []
    for c in captures:
        if is_main_account_url(c.url):
            out.append(c)
    return out


def build_account_locator(
    main_account_captures: list[Capture],
    platform_exhibit_counts: dict[str, int],
    platform_exhibit_filenames: dict[str, str],
    output_path: Path,
    case_name: str,
) -> list[AccountEntry]:
    """
    Write the Account Locator docx. Returns the entries it wrote (for manifest).

    `platform_exhibit_counts` maps platform_id -> number of exhibits in that
    platform's doc (0 if no posts). `platform_exhibit_filenames` maps
    platform_id -> the doc filename like "03_TikTok_Exhibits.pdf".
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"Account Locator — {case_name}", level=1)
    doc.add_paragraph(
        "Main account profile/landing-page captures, grouped by platform. "
        "Cross-references list the per-platform exhibit document where post "
        "captures from that account are filed."
    )

    by_platform: dict[str, list[Capture]] = {}
    for c in main_account_captures:
        plat = classify(c.url)
        by_platform.setdefault(plat, []).append(c)

    entries: list[AccountEntry] = []

    for platform in PLATFORM_ORDER:
        accounts = by_platform.get(platform, [])
        if not accounts:
            continue

        doc.add_heading(PLATFORM_DISPLAY_NAMES[platform], level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Handle"
        hdr[1].text = "Profile URL"
        hdr[2].text = "First Capture Date"
        hdr[3].text = "Cross-Reference"

        accounts_sorted = sorted(accounts, key=lambda c: c.capture_date_raw)
        for c in accounts_sorted:
            handle = extract_handle(c.url, platform)
            count = platform_exhibit_counts.get(platform, 0)
            fname = platform_exhibit_filenames.get(platform)
            if count and fname:
                xref = f"See {fname}, Exhibits 1–{count}"
            else:
                xref = "—"

            row = table.add_row().cells
            row[0].text = handle or "(unknown)"
            row[1].text = c.url
            row[2].text = c.capture_date_raw
            row[3].text = xref

            entries.append(AccountEntry(
                platform=platform,
                handle=handle,
                url=c.url,
                capture_date_raw=c.capture_date_raw,
                exhibit_doc_filename=fname,
                exhibit_count=count,
            ))

        doc.add_paragraph("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return entries
